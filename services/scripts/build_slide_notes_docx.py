#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_DEFAULT = Path("/Users/islam/IdeaProjects/RPS-limiter/services/Пояснения_к_слайдам_предзащиты.txt")
OUTPUT_DEFAULT = Path("/Users/islam/IdeaProjects/RPS-limiter/services/Пояснения_к_слайдам_предзащиты.docx")


SUBHEADINGS = {
    "Что здесь происходит",
    "Что говорить комиссии",
    "Возможные вопросы комиссии и ответы",
}


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_doc(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped == lines[0].strip():
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(15)
            continue

        if stripped.startswith("Слайд "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(13)
            continue

        if stripped in SUBHEADINGS or stripped.startswith("Вопрос "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(stripped)
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            continue

        if stripped == "Ответ.":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(stripped)
            run.bold = True
            run.italic = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        if stripped.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        run = paragraph.add_run(stripped)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    doc.save(output)
    print(f"Saved: {output}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a DOCX with explanations for presentation slides.")
    parser.add_argument("--input", type=Path, default=SOURCE_DEFAULT)
    parser.add_argument("--output", type=Path, default=OUTPUT_DEFAULT)
    args = parser.parse_args()

    build_doc(args.input, args.output)


if __name__ == "__main__":
    main()
