#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import re
import urllib.error
import urllib.parse
import urllib.request
from pathlib import Path
from typing import List, Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path("/Users/islam/IdeaProjects/RPS-limiter/services")
SRC = ROOT / "Диссертация.txt"
OUT = ROOT / "Диссертация_регламент.docx"
IMG_CACHE = ROOT / ".cache_vkr_images"


STRUCTURAL_HEADERS = {
    "Введение",
    "Заключение",
    "Список использованных источников",
    "Список иллюстративного материала",
    "Термины и определения",
    "Перечень сокращений и условных обозначений",
    "СОДЕРЖАНИЕ",
}


def normalize_quotes(text: str) -> str:
    # Normalize straight/curly quotes into Russian guillemets.
    value = (
        text.replace("“", "\"")
        .replace("”", "\"")
        .replace("„", "\"")
        .replace("‘", "'")
        .replace("’", "'")
    )
    # Repeat replacement to handle multiple quoted fragments in one line.
    prev = None
    while prev != value:
        prev = value
        value = re.sub(r"\"([^\"]+)\"", r"«\1»", value)
        value = re.sub(r"'([^']+)'", r"«\1»", value)
    return value


def set_run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    lang = run._element.rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.rPr.append(lang)
    lang.set(qn("w:val"), "ru-RU")


def add_runs_with_backtick_bold(paragraph, text: str, size: int, default_bold: bool = False):
    value = normalize_quotes(text)
    parts = re.split(r"(`[^`]+`)", value)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=default_bold)


def set_paragraph_base(paragraph, indent=True):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    text_elem = OxmlElement("w:t")
    text_elem.text = "1"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text_elem)
    run._r.append(fld_char_end)


def add_toc_field(paragraph):
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление будет сформировано в Word после обновления поля (F9)."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    set_run_font(run, size=14, bold=False)


def parse_blocks(lines: List[str]):
    # Use first occurrences for terms/abbr and body from the second "Введение"
    idx_terms = lines.index("Термины и определения")
    idx_abbr = lines.index("Перечень сокращений и условных обозначений")
    idx_toc = lines.index("СОДЕРЖАНИЕ")

    intro_positions = [i for i, line in enumerate(lines) if line == "Введение"]
    if len(intro_positions) < 2:
        raise ValueError("Ожидалось минимум два заголовка 'Введение' (в оглавлении и тексте).")

    idx_intro_text = intro_positions[1]
    idx_sources = max(i for i, line in enumerate(lines) if line == "Список использованных источников")
    idx_illustr = max(i for i, line in enumerate(lines) if line == "Список иллюстративного материала")

    topic = lines[0].strip() if lines and lines[0].startswith("Тема:") else ""

    terms = [x for x in lines[idx_terms + 1 : idx_abbr] if x.strip()]
    abbr = [x for x in lines[idx_abbr + 1 : idx_toc] if x.strip()]
    body = lines[idx_intro_text:idx_sources]
    sources = [x for x in lines[idx_sources + 1 : idx_illustr] if x.strip()]
    illustr = [x for x in lines[idx_illustr + 1 :] if x.strip()]

    return topic, terms, abbr, body, sources, illustr


def is_chapter(line: str) -> bool:
    return bool(re.match(r"^Глава\s+\d+\.", line))


def is_subheading(line: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\s", line))


def is_table_caption(line: str) -> bool:
    return bool(re.match(r"^Таблица\s+\d+\s+–\s+", line))


def is_figure_caption(line: str) -> bool:
    return bool(re.match(r"^Рисунок\s+\d+\s+–\s+", line))


def parse_markdown_row(row: str) -> List[str]:
    # | col1 | col2 |
    parts = [p.strip() for p in row.strip().strip("|").split("|")]
    return parts


def add_heading(document: Document, text: str, level: int, structural: bool = False):
    p = document.add_paragraph(style=f"Heading {level}")
    heading_text = normalize_quotes(text.upper() if structural else text)
    run = p.add_run(heading_text)
    set_run_font(run, size=14, bold=True)

    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    if structural:
        pf.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_plain_paragraph(document: Document, text: str, indent=True, center=False):
    p = document.add_paragraph()
    add_runs_with_backtick_bold(p, text, size=14, default_bold=False)
    set_paragraph_base(p, indent=indent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table_from_block(document: Document, block: List[str]):
    rows = [parse_markdown_row(line) for line in block if line.strip()]
    # drop markdown separator rows like |---|---|
    filtered = []
    for r in rows:
        if r and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in r):
            continue
        filtered.append(r)

    if not filtered:
        return

    cols = max(len(r) for r in filtered)
    table = document.add_table(rows=len(filtered), cols=cols)
    table.style = "Table Grid"

    for i, r in enumerate(filtered):
        for j in range(cols):
            cell_text = r[j] if j < len(r) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_backtick_bold(p, cell_text, size=12, default_bold=(i == 0))
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                pf.first_line_indent = Cm(0)


def insert_page_break(document: Document):
    document.add_page_break()


def build_placeholder_png(ref: str) -> Path:
    IMG_CACHE.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(ref.encode("utf-8")).hexdigest()[:16]
    out = IMG_CACHE / f"{digest}_placeholder.png"
    if out.exists():
        return out

    parsed = urllib.parse.urlparse(ref)
    text = urllib.parse.parse_qs(parsed.query).get("text", ["Иллюстрация"])[0]
    title = text.replace("+", " ")

    fig, ax = plt.subplots(figsize=(14, 7), dpi=100)
    fig.patch.set_facecolor("#f2f2f2")
    ax.set_facecolor("#f2f2f2")
    ax.axis("off")
    ax.text(
        0.5,
        0.5,
        title,
        ha="center",
        va="center",
        fontsize=26,
        wrap=True,
        family="DejaVu Sans",
        color="#202020",
    )
    fig.savefig(out, format="png", dpi=100, bbox_inches="tight", pad_inches=0.2)
    plt.close(fig)
    return out


def resolve_image_path(raw_ref: str) -> Optional[Path]:
    ref = raw_ref.strip()
    if not ref:
        return None

    if ref.startswith("http://") or ref.startswith("https://"):
        if "placehold.co" in ref:
            try:
                return build_placeholder_png(ref)
            except Exception:
                return None

        IMG_CACHE.mkdir(parents=True, exist_ok=True)
        digest = hashlib.sha256(ref.encode("utf-8")).hexdigest()[:16]
        local = IMG_CACHE / f"{digest}.png"
        if local.exists():
            return local
        try:
            safe_url = urllib.parse.quote(ref, safe=":/?&=+,%.-_~#")
            with urllib.request.urlopen(safe_url, timeout=20) as response:
                data = response.read()
            if data.lstrip().startswith(b"<svg") or data.lstrip().startswith(b"<!DOCTYPE html"):
                return None
            local.write_bytes(data)
            return local
        except (urllib.error.URLError, TimeoutError, OSError):
            return None

    path = Path(ref)
    if path.is_absolute() and path.exists():
        return path
    local = ROOT / ref
    if local.exists():
        return local
    return None


def add_image_from_reference(document: Document, raw_ref: str):
    image_path = resolve_image_path(raw_ref)
    if image_path is None:
        add_plain_paragraph(document, f"Ссылка на изображение: {raw_ref}", indent=False)
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    # Available width with margins is 165 mm. Keep a small visual margin.
    try:
        run.add_picture(str(image_path), width=Cm(16))
    except Exception:
        add_plain_paragraph(document, f"Ссылка на изображение: {raw_ref}", indent=False)


def build_doc(topic: str, terms: List[str], abbr: List[str], body: List[str], sources: List[str], illustr: List[str]):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(15)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    # Heading styles map to TOC levels
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h2.font.size = Pt(14)
    h2.font.bold = True

    # Footer page number centered
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_p)
    for run in footer_p.runs:
        set_run_font(run, size=14, bold=False)

    # Optional topic line
    if topic:
        p = doc.add_paragraph()
        run = p.add_run(topic)
        set_run_font(run, size=14, bold=False)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_base(p, indent=False)
        insert_page_break(doc)

    # CONTENTS
    add_heading(doc, "СОДЕРЖАНИЕ", level=1, structural=True)
    toc_p = doc.add_paragraph()
    set_paragraph_base(toc_p, indent=False)
    add_toc_field(toc_p)

    # Terms
    insert_page_break(doc)
    add_heading(doc, "Термины и определения", level=1, structural=True)
    for line in terms:
        add_plain_paragraph(doc, line)

    # Abbreviations
    insert_page_break(doc)
    add_heading(doc, "Перечень сокращений и условных обозначений", level=1, structural=True)
    for line in abbr:
        add_plain_paragraph(doc, line)

    # Main body: from Introduction to before sources
    i = 0
    first_main_heading_seen = False
    while i < len(body):
        line = body[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Structural headings inside body
        if stripped in {"Введение", "Заключение"}:
            insert_page_break(doc)
            add_heading(doc, stripped, level=1, structural=True)
            first_main_heading_seen = True
            i += 1
            continue

        if is_chapter(stripped):
            insert_page_break(doc)
            add_heading(doc, stripped, level=1, structural=False)
            first_main_heading_seen = True
            i += 1
            continue

        if is_subheading(stripped):
            add_heading(doc, stripped, level=2, structural=False)
            i += 1
            continue

        if stripped.startswith("Выводы по главе"):
            add_heading(doc, stripped, level=2, structural=False)
            i += 1
            continue

        if is_table_caption(stripped):
            p = add_plain_paragraph(doc, stripped, indent=False)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Capture markdown table lines following caption
            j = i + 1
            table_block: List[str] = []
            while j < len(body) and body[j].strip().startswith("|"):
                table_block.append(body[j].strip())
                j += 1
            add_table_from_block(doc, table_block)
            i = j
            continue

        if is_figure_caption(stripped):
            p = add_plain_paragraph(doc, stripped, indent=False, center=True)
            for run in p.runs:
                run.bold = False
            i += 1
            continue

        if stripped.startswith("Ссылка на изображение:"):
            image_ref = stripped.split(":", 1)[1].strip() if ":" in stripped else ""
            add_image_from_reference(doc, image_ref)
            i += 1
            continue

        add_plain_paragraph(doc, stripped)
        i += 1

    # Sources
    insert_page_break(doc)
    add_heading(doc, "Список использованных источников", level=1, structural=True)
    for line in sources:
        add_plain_paragraph(doc, line)

    # Illustrative list
    insert_page_break(doc)
    add_heading(doc, "Список иллюстративного материала", level=1, structural=True)
    for line in illustr:
        add_plain_paragraph(doc, line)

    doc.save(OUT)


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    topic, terms, abbr, body, sources, illustr = parse_blocks(lines)
    build_doc(topic, terms, abbr, body, sources, illustr)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
