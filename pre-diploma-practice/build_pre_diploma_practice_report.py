from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
SRC = BASE / 'Отчёт_производственная_преддипломная_практика_RPS_лимитирование.txt'
OUT = BASE / 'Отчёт_производственная_преддипломная_практика_RPS_лимитирование.docx'

text = SRC.read_text(encoding='utf-8')

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(10)


def set_paragraph_format(p, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if first_line:
        pf.first_line_indent = Cm(1.25)


def add_paragraph(doc, line):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(12)
    return p


def add_heading(doc, line, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level == 0 else 8)
    pf.space_after = Pt(12 if level == 0 else 6)
    if level > 0:
        pf.first_line_indent = Cm(1.25)
    run = p.add_run(line)
    run.bold = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(14 if level in (0, 1) else 12)
    return p


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, value.strip(), bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, 'D9EAF7')
    doc.add_paragraph()
    return table


def add_image(doc, image_name):
    image_path = BASE / image_name
    if not image_path.exists():
        raise FileNotFoundError(f'Image not found: {image_path}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(16.0))
    return p


def title_paragraph(doc, text='', *, bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=0, space_after=0, line_spacing=1.15, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.first_line_indent = Cm(0)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    run = p.add_run(text if text else ' ')
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    return p


def add_title_page(doc):
    title_paragraph(doc, 'Министерство науки и высшего образования Российской Федерации', bold=True)
    title_paragraph(doc, 'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ', bold=True)
    title_paragraph(doc, 'ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ', bold=True)
    title_paragraph(doc, '«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ', bold=True)
    title_paragraph(doc, 'ИТМО»', bold=True)
    title_paragraph(doc, '(Университет ИТМО)', bold=True, space_after=10)
    title_paragraph(doc, 'Факультет программной инженерии и компьютерной техники', bold=True, space_after=6)
    title_paragraph(doc, 'Образовательная программа: 09.04.01 Компьютерные системы и технологии',
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    title_paragraph(doc, 'Направление подготовки (специальность): Программная инженерия',
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=24)
    title_paragraph(doc, 'О Т Ч Е Т', bold=True, size=16, space_after=4)
    title_paragraph(doc, 'о производственной, преддипломной практике', size=13, space_after=8)
    title_paragraph(
        doc,
        'Тема задания: «Адаптивное управление нагрузкой в распределенной среде на основе интеллектуального анализа трафика»',
        size=13,
        space_after=16,
    )
    title_paragraph(doc, 'Обучающийся Юнусов Ислам Рамилевич, № P4219',
                    align=WD_ALIGN_PARAGRAPH.LEFT, size=14, space_after=6)
    title_paragraph(doc, 'Руководитель практики от университета: Платонов Алексей Владимирович',
                    align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=8)
    title_paragraph(doc, 'Практика пройдена с оценкой ____',
                    align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=7.8, size=14, space_after=8)
    title_paragraph(doc, 'Подписи членов комиссии:',
                    align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=7.8, size=14, space_after=8)
    for _ in range(3):
        title_paragraph(doc, '______________ Ф.И.О.',
                        align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=9.5, size=13)
        title_paragraph(doc, '(подпись)',
                        align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=9.6, size=12, space_after=4)
    title_paragraph(doc, 'Дата 29.05.2026',
                    align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=7.8, size=13, space_after=18)
    title_paragraph(doc, 'Санкт-Петербург', size=14)
    title_paragraph(doc, '2026', size=14)

# Document setup
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)
section.different_first_page_header_footer = True

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
styles['Normal'].font.size = Pt(12)

settings = doc.settings.element
auto_hyphenation = OxmlElement('w:autoHyphenation')
auto_hyphenation.set(qn('w:val'), 'true')
settings.append(auto_hyphenation)

# Split title page / rest
pages = text.split('\n===PAGE===\n')

add_title_page(doc)
doc.add_page_break()

# Body pages
for page_idx, page in enumerate(pages[1:], start=1):
    if page_idx > 1:
        doc.add_page_break()
    lines = page.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line == '[PAGEBREAK]':
            doc.add_page_break()
            i += 1
            continue
        image_match = re.match(r'^\[IMAGE:\s*(.+?)\s*\]$', line)
        if image_match:
            add_image(doc, image_match.group(1))
            i += 1
            continue
        if page_idx == 1 and line != 'СОДЕРЖАНИЕ':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(3)
            pf.first_line_indent = Cm(0)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(12)
            i += 1
            continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                cells = [c.strip() for c in row_line.strip('|').split('|')]
                # Skip markdown separator rows.
                if not all(set(c.replace(' ', '')) <= {'-'} for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        if line in {'СОДЕРЖАНИЕ', 'ВВЕДЕНИЕ', 'ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ', 'ЗАКЛЮЧЕНИЕ', 'ПРИМЕЧАНИЯ И ПРИЛОЖЕНИЯ'}:
            add_heading(doc, line, 0)
        elif line.startswith('Приложение '):
            add_heading(doc, line, 1)
        elif re.match(r'^[1-6]\. ', line):
            add_heading(doc, line, 1)
        elif re.match(r'^[1-6]\.[1-9]\. ', line):
            add_heading(doc, line, 2)
        elif line.startswith('Таблица '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(6)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(12)
        else:
            add_paragraph(doc, line)
        i += 1

# Footer page numbers
for sec in doc.sections:
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    field.append(r)
    p._p.append(field)

doc.save(OUT)
print(OUT)
